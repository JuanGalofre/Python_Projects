import pandas
from clase import Persona
from docx import Document
from docx.shared import Pt
df = pandas.read_excel("defaultexcel.xlsx")
numero_de_columnas= df.shape[0]
listaPersonas=[]

for i in range(numero_de_columnas):
    nombre=None
    cedula=None
    ciudad=None
    referencia=None
    fecha=None
    sexo=None
    inversion=None
    nombre=df.loc[i][2].strip()
    cedula=df.loc[i][4]
    ciudad=df.loc[i][6]
    referencia=df.loc[i][5]
    fecha=df.loc[i][0]
    inversion= str(df.loc[i][1])
    listaI=list(inversion)
    contador=1
    cc=0
    for i in listaI:
        if contador%3 ==0:
            listaI.insert(-contador+cc,".")
            cc-=1
        print(listaI)
        contador+=1
    while listaI[0] ==".":
        listaI.pop(0)
    inversion = "".join(listaI)
    persona=Persona(nombre,cedula,ciudad,referencia,fecha,sexo,inversion)
    listaPersonas.append(persona)
for i in listaPersonas:
    tamaño= Pt(12)
    fuente= 'Arial'
    document= Document("Sampletext.docx")
    parrafo = document.paragraphs[0]
    parrafo = parrafo.clear()
    run= parrafo.add_run(("Something ").upper())
    run.font.size = tamaño
    run.font.name = fuente
    run.font.bold = True
    runNombre= parrafo.add_run(str(i.darNombre()))
    runNombre.font.size = tamaño
    runNombre.font.name = fuente
    runNombre.font.bold = True
    runConector=parrafo.add_run(" and ")
    runConector.font.size = tamaño
    runConector.font.name = fuente
    runConector.font.bold = False
    run1= parrafo.add_run("Someoone")
    run1.font.size = tamaño
    run1.font.name = fuente
    run1.font.bold = True

    parrafo1 = document.paragraphs[19]
    parrafo1 = parrafo1.clear()
    runp1 = parrafo1.add_run("He ")
    runp1.font.size = tamaño
    runp1.font.name = fuente
    runNombre1 = parrafo1.add_run(str(i.darNombre()))
    runNombre1.font.size = tamaño
    runNombre1.font.name = fuente
    runNombre1.font.bold = True
    runp2 = parrafo1.add_run("does something with his ID: ")
    runp2.font.size = tamaño
    runp2.font.name = fuente
    runCedula = parrafo1.add_run(str(i.darCedula()))
    runCedula.font.size = tamaño
    runCedula.font.name = fuente
    runCedula.font.bold = True
    runp3= parrafo1.add_run(" something in")
    runp3.font.size = tamaño
    runp3.font.name = fuente
    runCiudad = parrafo1.add_run(str(i.darCiudad()))
    runCiudad.font.size = tamaño
    runCiudad.font.name = fuente

    parrafo2 = document.paragraphs[24]
    runp4= parrafo2.add_run("because ")
    runp4.font.size = tamaño
    runp4.font.name = fuente
    runp5= parrafo2.add_run(" some reason")
    runp5.font.size = tamaño
    runp5.font.name = fuente
    runp5.font.bold = True

    string="Some_Stuff"
    sinespacios=str(i.darNombre()).strip()
    recortes=sinespacios.split(" ")
    iniciales=""
    for j in recortes:
        iniciales = iniciales + str(j[0])
    string= string + iniciales +"_withReference_"
    referenciaNombre=str(i.darReferencia()).replace("REFERENCIA No. ","")
    string= string+referenciaNombre+".docx"
    document.save(string)
